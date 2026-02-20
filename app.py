import os
import pathlib
import re
import zlib
from datetime import datetime

from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from werkzeug.utils import secure_filename
from docx import Document

import db
from ai_processor import process_text_with_ai, run_project_action

load_dotenv(override=True)

APP_DIR = pathlib.Path(__file__).resolve().parent
UPLOAD_DIR = APP_DIR / "uploads"
OUTPUT_DIR = APP_DIR / "outputs"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

ALLOWED_EXT = {".docx", ".txt"}
MAX_CONTENT_LENGTH = 20 * 1024 * 1024  # 20MB

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev-secret")
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH


def allowed_file(filename: str) -> bool:
    ext = pathlib.Path(filename).suffix.lower()
    return ext in ALLOWED_EXT


# --------------------------
# PLANTUML SERVER
# --------------------------
PLANTUML_SERVER = os.getenv("PLANTUML_SERVER", "https://www.plantuml.com/plantuml")


def _plantuml_encode64(data: bytes) -> str:
    alphabet = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"
    res = []
    i = 0
    while i < len(data):
        b1 = data[i]; i += 1
        b2 = data[i] if i < len(data) else 0; i += 1
        b3 = data[i] if i < len(data) else 0; i += 1

        c1 = b1 >> 2
        c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
        c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
        c4 = b3 & 0x3F

        res.append(alphabet[c1 & 0x3F])
        res.append(alphabet[c2 & 0x3F])
        res.append(alphabet[c3 & 0x3F])
        res.append(alphabet[c4 & 0x3F])

    return "".join(res)


def _encode_plantuml_deflate(plantuml_text: str) -> str:
    data = plantuml_text.encode("utf-8")
    compressed = zlib.compress(data)[2:-4]  # zlib header/footer at
    return _plantuml_encode64(compressed)


def plantuml_image_url(plantuml_code: str, fmt: str = "png") -> str:
    encoded = _encode_plantuml_deflate(plantuml_code)
    return f"{PLANTUML_SERVER}/{fmt}/{encoded}"


# -------------------------
# PlantUML temizleme
# -------------------------
def sanitize_plantuml(text: str) -> str:
    try:
        if not text:
            return "```plantuml\n@startuml\n@enduml\n```"

        m = re.search(r"```plantuml\s*(.*?)```", text, flags=re.DOTALL | re.IGNORECASE)
        code = (m.group(1) if m else text).strip()

        cleaned_lines = []
        for ln in code.splitlines():
            ln2 = ln.rstrip()
            if ln2.endswith(","):
                ln2 = ln2[:-1].rstrip()
            cleaned_lines.append(ln2)

        code = "\n".join(cleaned_lines).strip()

        low = code.lower()
        if "@startuml" not in low:
            code = "@startuml\n" + code
        if "@enduml" not in low:
            code = code + "\n@enduml"

        return "```plantuml\n" + code.strip() + "\n```"
    except Exception:
        return "```plantuml\n" + (text or "").strip() + "\n```"


def extract_plantuml_code(output_text: str) -> str:
    """
    sanitize_plantuml sonrası ```plantuml ... ``` bloğunun içini döndürür
    """
    m = re.search(r"```plantuml\s*(.*?)```", output_text, flags=re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return output_text.strip()


# =========================
# 1) TEK ANA SAYFA: PROJECTS (index.html)
# =========================
@app.get("/")
def index():
    rows = db.list_projects(200)
    return render_template("index.html", rows=rows)


# =========================
# 2) DOWLAMD MODÜLÜ ENDPOINTS
# =========================

@app.get("/download/<int:file_id>")
def download(file_id: int):
    row = db.get_file(file_id)
    if not row or row["status"] != "DONE" or not row.get("output_path"):
        flash("Bu dosya henüz indirilebilir değil.", "error")
        return redirect(url_for("index"))
    return send_file(row["output_path"], as_attachment=True)


# =========================
# 3) PROJECT CREATE
# =========================
@app.post("/projects/create")
def projects_create():
    data = {
        "title": (request.form.get("title") or "").strip(),
        "domain": (request.form.get("domain") or "").strip(),
        "primary_entity": (request.form.get("primary_entity") or "").strip(),
        "constraints_text": (request.form.get("constraints_text") or "").strip(),
        "advanced_feature": (request.form.get("advanced_feature") or "").strip(),
        "security_access": (request.form.get("security_access") or "").strip(),
        "reporting_requirement": (request.form.get("reporting_requirement") or "").strip(),
        "common_tasks": (request.form.get("common_tasks") or "").strip(),
    }

    if not data["title"] or not data["domain"] or not data["primary_entity"]:
        flash("Title, Domain ve Primary Entity zorunlu.", "error")
        return redirect(url_for("index"))

    existing = db.get_project_by_title(data["title"])
    if existing:
        flash("Bu başlıkta bir proje zaten var. Lütfen farklı bir isim gir.", "error")
        return redirect(url_for("index"))

    try:
        project_id = db.insert_project(data)
    except Exception as e:
        if "Duplicate entry" in str(e) or "uq_projects_title" in str(e):
            flash("Bu başlıkta bir proje zaten var. Lütfen farklı bir isim gir.", "error")
            return redirect(url_for("index"))
        raise

    flash(f"Project oluşturuldu (ID={project_id}).", "ok")
    return redirect(url_for("project_detail", project_id=project_id))


# =========================
# 4) CREATE + ALL GENERATE DOCX
# =========================
ALL_ACTIONS = [
    ("business_rules", "Business Rules"),
    ("er_tables", "ER Tables"),
    ("missing_rules", "Missing Rules"),
    ("normalization", "Normalization (0NF → 3NF)"),
    ("er_plantuml", "ER Diagram (PlantUML)"),
    ("sql_script", "SQL Script"),
    ("report", "Report Queries"),
]


def _is_md_table(text: str) -> bool:
    if not text:
        return False
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if len(lines) < 3:
        return False

    if not lines[0].startswith("|") or lines[0].count("|") < 3:
        return False

    sep = lines[1]
    if not sep.startswith("|") or sep.count("|") < 3:
        return False

    allowed = set("|:- ")
    if any(ch not in allowed and ch != "-" for ch in sep):
        return False
    if "-" not in sep:
        return False

    if not lines[2].startswith("|"):
        return False

    return True


def _parse_md_table(text: str):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    table_lines = []
    for ln in lines:
        if ln.startswith("|"):
            table_lines.append(ln)
        elif table_lines:
            break

    if len(table_lines) < 3:
        return None

    def split_row(ln: str):
        core = ln.strip().strip("|")
        return [c.strip() for c in core.split("|")]

    header = split_row(table_lines[0])
    rows = [split_row(ln) for ln in table_lines[2:] if ln.startswith("|")]

    col_count = len(header)
    fixed_rows = []
    for r in rows:
        if len(r) < col_count:
            r = r + [""] * (col_count - len(r))
        elif len(r) > col_count:
            r = r[:col_count]
        fixed_rows.append(r)

    return header, fixed_rows


def _docx_add_md_table(doc: Document, title: str, md_table_text: str) -> bool:
    parsed = _parse_md_table(md_table_text)
    if not parsed:
        return False

    header, rows = parsed

    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h

    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = val

    doc.add_paragraph("")
    return True


def _docx_add_block(doc: Document, title: str, content: str, allow_table: bool = True):
    if allow_table and _is_md_table(content):
        if _docx_add_md_table(doc, title, content):
            return

    doc.add_heading(title, level=2)
    for line in (content or "").splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph("")


@app.post("/projects/create_and_generate")
def projects_create_and_generate():
    data = {
        "title": (request.form.get("title") or "").strip(),
        "domain": (request.form.get("domain") or "").strip(),
        "primary_entity": (request.form.get("primary_entity") or "").strip(),
        "constraints_text": (request.form.get("constraints_text") or "").strip(),
        "advanced_feature": (request.form.get("advanced_feature") or "").strip(),
        "security_access": (request.form.get("security_access") or "").strip(),
        "reporting_requirement": (request.form.get("reporting_requirement") or "").strip(),
        "common_tasks": (request.form.get("common_tasks") or "").strip(),
    }

    if not data["title"] or not data["domain"] or not data["primary_entity"]:
        flash("Title, Domain ve Primary Entity zorunlu.", "error")
        return redirect(url_for("index"))

    existing = db.get_project_by_title(data["title"])
    if existing:
        flash("Bu başlıkta bir proje zaten var. Lütfen farklı bir isim gir.", "error")
        return redirect(url_for("index"))

    try:
        project_id = db.insert_project(data)
    except Exception as e:
        if "Duplicate entry" in str(e) or "uq_projects_title" in str(e):
            flash("Bu başlıkta bir proje zaten var. Lütfen farklı bir isim gir.", "error")
            return redirect(url_for("index"))
        raise

    p = db.get_project(project_id)

    doc = Document()
    doc.add_heading(p["title"], level=1)
    doc.add_paragraph(f"Domain: {p['domain']}")
    doc.add_paragraph(f"Primary Entity: {p['primary_entity']}")
    doc.add_paragraph("")

    failures = []

    for action_key, section_title in ALL_ACTIONS:
        try:
            prompt_text, output_text, model_used = run_project_action(p, action_key, temperature=0.2)

        
            if action_key == "er_plantuml":
                output_text = sanitize_plantuml(output_text)
                _docx_add_block(doc, section_title, output_text, allow_table=False)
            else:
                _docx_add_block(doc, section_title, output_text, allow_table=True)

            db.insert_project_output(
                project_id=project_id,
                action_key=action_key,
                prompt_text=prompt_text,
                output_text=output_text,
                model=model_used,
                temperature=0.2,
            )

        except Exception as e:
            failures.append(f"{action_key}: {e}")
            _docx_add_block(doc, f"{section_title} (HATA)", str(e), allow_table=False)

    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / f"project_{project_id}_all_{ts}.docx"
    doc.save(str(out_path))

    file_id = db.insert_file(
        original_name=f"PROJECT_{project_id}_ALL_OUTPUTS.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        input_path="",
    )
    db.set_status(file_id, "DONE", error_message=None, output_path=str(out_path))

    if failures:
        flash("Hepsi üretildi ama bazı bölümlerde hata oldu: " + " | ".join(failures), "error")
    else:
        flash("Hepsi üretildi. DOCX hazır!", "ok")

    return redirect(url_for("download", file_id=file_id))


# =========================
# 5) PROJECT DETAIL
# =========================
@app.get("/project/<int:project_id>")
def project_detail(project_id: int):
    p = db.get_project(project_id)
    if not p:
        flash("Project bulunamadı.", "error")
        return redirect(url_for("index"))

    latest = {
        "business_rules": db.get_latest_project_output(project_id, "business_rules"),
        "er_tables": db.get_latest_project_output(project_id, "er_tables"),
        "missing_rules": db.get_latest_project_output(project_id, "missing_rules"),
        "normalization": db.get_latest_project_output(project_id, "normalization"),
        "er_plantuml": db.get_latest_project_output(project_id, "er_plantuml"),
        "sql_script": db.get_latest_project_output(project_id, "sql_script"),
        "report": db.get_latest_project_output(project_id, "report"),
    }

    return render_template("project_detail.html", p=p, latest=latest)


# =========================
# 6) RUN ACTION 
# =========================
@app.post("/project/<int:project_id>/run/<action_key>")
def project_run(project_id: int, action_key: str):
    p = db.get_project(project_id)
    if not p:
        flash("Project bulunamadı.", "error")
        return redirect(url_for("index"))
    

    try:
        prompt_text, output_text, model_used = run_project_action(p, action_key, temperature=0.2)

        img_url = None
        if action_key == "er_plantuml":
            output_text = sanitize_plantuml(output_text)
            code = extract_plantuml_code(output_text)
            img_url = plantuml_image_url(code, fmt="svg")

        out_id = db.insert_project_output(
            project_id=project_id,
            action_key=action_key,
            prompt_text=prompt_text,
            output_text=output_text,
            model=model_used,
            temperature=0.2,
        )

        return render_template(
            "index_yeni.html",
            p=p,
            action_key=action_key,
            output_text=output_text,
            prompt_text=prompt_text,
            model=model_used,
            out_id=out_id,
            img_url=img_url,   
        )

    except Exception as e:
        return render_template(
            "index_yeni.html",
            p=p,
            action_key=action_key,
            error=str(e),
        )


if __name__ == "__main__":
    app.run(debug=True)
