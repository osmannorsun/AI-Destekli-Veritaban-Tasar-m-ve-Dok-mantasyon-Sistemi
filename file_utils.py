import pathlib
from docx import Document

def extract_text(path: pathlib.Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        doc = Document(str(path))
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        return "\n".join(parts).strip()

    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore").strip()

    raise ValueError("Desteklenmeyen dosya türü. (.docx, .txt)")

def write_docx_from_text(text: str, out_path: pathlib.Path):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(out_path))
